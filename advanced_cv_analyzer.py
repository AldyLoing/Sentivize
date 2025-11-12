"""
Advanced CV/Resume Analyzer - Deep Parsing & Contextual Understanding
Sistem AI yang benar-benar memahami isi CV dan dapat menilai kecocokan secara kontekstual
"""

import pandas as pd
from typing import Dict, List, Optional, Tuple, Any
import PyPDF2
import docx
import re
from pathlib import Path
from dataclasses import dataclass
from datetime import datetime
import json

from advanced_ai_core import (
    AdvancedNLPEngine, 
    ContextualReasoningEngine,
    get_advanced_nlp_engine,
    get_reasoning_engine,
    AnalysisInsight
)
import config

# Import enhanced parser
try:
    from enhanced_cv_parser import EnhancedCVParser
    ENHANCED_PARSER_AVAILABLE = True
except ImportError:
    ENHANCED_PARSER_AVAILABLE = False


@dataclass
class CVSection:
    """Struktur data untuk section CV"""
    section_type: str
    content: str
    structured_data: Dict[str, Any]
    importance_score: float


@dataclass
class WorkExperience:
    """Struktur data pengalaman kerja"""
    company: str
    position: str
    duration: str
    start_date: Optional[str]
    end_date: Optional[str]
    years: float
    responsibilities: List[str]
    achievements: List[str]
    technologies: List[str]
    relevance_score: float


@dataclass
class Education:
    """Struktur data pendidikan"""
    institution: str
    degree: str
    field_of_study: str
    graduation_year: Optional[int]
    gpa: Optional[float]
    achievements: List[str]


@dataclass
class CVProfile:
    """Profile lengkap dari CV"""
    candidate_name: str
    contact_info: Dict[str, str]
    professional_summary: str
    work_experiences: List[WorkExperience]
    education: List[Education]
    skills: Dict[str, List[str]]  # technical, soft, languages, tools
    certifications: List[str]
    projects: List[Dict[str, str]]
    awards: List[str]
    interests: List[str]
    total_experience_years: float
    seniority_level: str


@dataclass
@dataclass
class CVAnalysisResult:
    """Hasil analisis CV yang komprehensif"""
    cv_profile: CVProfile
    relevance_score: float
    relevance_breakdown: Dict[str, float]
    relevance_reasoning: str  # Added: reasoning untuk relevansi
    strengths: List[AnalysisInsight]
    potential_areas: List[AnalysisInsight]
    fit_analysis: Dict[str, Any]
    professional_assessment: str
    soft_skills_assessment: Dict[str, float]
    overall_recommendation: str
    confidence_score: float


class AdvancedCVParser:
    """
    Advanced CV Parser dengan deep understanding
    """
    
    def __init__(self):
        self.section_headers = {
            'summary': ['summary', 'profile', 'about', 'objective', 'ringkasan', 'profil', 'tentang'],
            'experience': ['experience', 'work history', 'employment', 'pengalaman', 'riwayat kerja', 'pekerjaan'],
            'education': ['education', 'academic', 'pendidikan', 'akademik', 'riwayat pendidikan'],
            'skills': ['skills', 'competencies', 'expertise', 'kemampuan', 'keahlian', 'kompetensi'],
            'certifications': ['certifications', 'certificates', 'sertifikat', 'sertifikasi', 'lisensi'],
            'projects': ['projects', 'portfolio', 'proyek', 'portofolio'],
            'awards': ['awards', 'achievements', 'honors', 'penghargaan', 'prestasi'],
            'interests': ['interests', 'hobbies', 'minat', 'hobi']
        }
        
        # Initialize enhanced parser if available
        self.enhanced_parser = EnhancedCVParser() if ENHANCED_PARSER_AVAILABLE else None
    
    def parse_with_enhanced_parser(self, text: str) -> Dict:
        """
        Use enhanced parser untuk ekstraksi maksimal
        Returns all extracted information in structured format
        """
        if not self.enhanced_parser:
            return None
        
        try:
            extracted = self.enhanced_parser.extract_all_info(text)
            return extracted
        except Exception as e:
            print(f"Enhanced parser error: {str(e)}")
            return None
    
    def extract_text_from_pdf(self, pdf_file) -> str:
        """Extract text dari PDF dengan better formatting"""
        try:
            # Try PyPDF2 first
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            text_parts = []
            
            for page in pdf_reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            
            full_text = "\n".join(text_parts)
            
            # If PyPDF2 fails or gives poor results, try pdfplumber
            if len(full_text.strip()) < 100:
                try:
                    import pdfplumber
                    pdf_file.seek(0)
                    with pdfplumber.open(pdf_file) as pdf:
                        text_parts = []
                        for page in pdf.pages:
                            text = page.extract_text()
                            if text:
                                text_parts.append(text)
                        full_text = "\n".join(text_parts)
                except ImportError:
                    pass
            
            return full_text
            
        except Exception as e:
            print(f"Error extracting PDF: {str(e)}")
            return ""
    
    def extract_text_from_docx(self, docx_file) -> str:
        """Extract text dari DOCX dengan better structure"""
        try:
            doc = docx.Document(docx_file)
            text_parts = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        text_parts.append(row_text)
            
            return "\n".join(text_parts)
            
        except Exception as e:
            print(f"Error extracting DOCX: {str(e)}")
            return ""
    
    def parse_cv_structure(self, text: str) -> Dict[str, CVSection]:
        """
        Parse CV ke dalam sections yang terstruktur
        
        Args:
            text: Raw CV text
            
        Returns:
            Dict of CVSection objects
        """
        sections = {}
        lines = text.split('\n')
        current_section = 'unknown'
        current_content = []
        
        for line in lines:
            line_lower = line.lower().strip()
            
            # Check if line is a section header
            section_found = None
            for section_type, headers in self.section_headers.items():
                for header in headers:
                    if header in line_lower and len(line_lower) < 50:
                        section_found = section_type
                        break
                if section_found:
                    break
            
            if section_found:
                # Save previous section
                if current_content:
                    content_text = "\n".join(current_content)
                    if current_section not in sections:
                        sections[current_section] = CVSection(
                            section_type=current_section,
                            content=content_text,
                            structured_data={},
                            importance_score=0.5
                        )
                
                # Start new section
                current_section = section_found
                current_content = []
            else:
                if line.strip():
                    current_content.append(line.strip())
        
        # Save last section
        if current_content:
            content_text = "\n".join(current_content)
            if current_section not in sections:
                sections[current_section] = CVSection(
                    section_type=current_section,
                    content=content_text,
                    structured_data={},
                    importance_score=0.5
                )
        
        return sections
    
    def extract_contact_info(self, text: str) -> Dict[str, str]:
        """Extract kontak informasi"""
        contact = {
            'email': '',
            'phone': '',
            'linkedin': '',
            'github': '',
            'location': ''
        }
        
        # Email
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        if emails:
            contact['email'] = emails[0]
        
        # Phone
        phone_pattern = r'(\+?\d{1,3}[-.\s]?)?\(?\d{2,4}\)?[-.\s]?\d{3,4}[-.\s]?\d{4}'
        phones = re.findall(phone_pattern, text)
        if phones:
            phone_str = phones[0] if isinstance(phones[0], str) else ''.join(phones[0])
            contact['phone'] = phone_str
        
        # LinkedIn
        linkedin_pattern = r'linkedin\.com/in/[\w-]+'
        linkedin_matches = re.findall(linkedin_pattern, text, re.IGNORECASE)
        if linkedin_matches:
            contact['linkedin'] = linkedin_matches[0]
        
        # GitHub
        github_pattern = r'github\.com/[\w-]+'
        github_matches = re.findall(github_pattern, text, re.IGNORECASE)
        if github_matches:
            contact['github'] = github_matches[0]
        
        return contact
    
    def parse_work_experiences(self, experience_text: str) -> List[WorkExperience]:
        """
        Parse pengalaman kerja dengan pattern matching yang lebih canggih
        """
        experiences = []
        
        # Multiple splitting strategies
        # Strategy 1: Split by date patterns at start of line
        entries = re.split(r'\n(?=\d{4}\s*[-–]\s*(?:\d{4}|Present|Current|Now|Sekarang))', experience_text)
        
        # Strategy 2: If no split, try company patterns
        if len(entries) <= 1:
            entries = re.split(r'\n(?=[A-Z][A-Za-z\s&.,Ltd]+(?:\s*\||$))', experience_text)
        
        # Strategy 3: Try position patterns
        if len(entries) <= 1:
            position_keywords = ['engineer', 'developer', 'manager', 'analyst', 'designer', 'consultant', 'specialist', 'coordinator', 'director']
            pattern = r'\n(?=(?:' + '|'.join(position_keywords) + r'))'
            entries = re.split(pattern, experience_text, flags=re.IGNORECASE)
        
        for entry in entries:
            if len(entry.strip()) < 20:
                continue
            
            lines = [l.strip() for l in entry.split('\n') if l.strip()]
            if not lines:
                continue
            
            company = ""
            position = ""
            duration_str = ""
            start_date = None
            end_date = None
            years = 0.0
            
            # Strategy 1: Look for date pattern first (most reliable anchor)
            for i, line in enumerate(lines):
                if re.search(r'\d{4}\s*[-–]\s*(?:\d{4}|Present|Current|Now|Sekarang)', line, re.IGNORECASE):
                    duration_str = line
                    
                    # Position usually before date, company might be before or after
                    if i > 0:
                        position = lines[i-1]
                    if i > 1:
                        company = lines[i-2]
                    elif i == 0 and len(lines) > 1:
                        # Date first, then position/company
                        position = lines[1] if len(lines) > 1 else ""
                        company = lines[2] if len(lines) > 2 else ""
                    break
            
            # Strategy 2: If no date found, look for pipe separator pattern
            if not duration_str:
                for line in lines:
                    if '|' in line:
                        parts = [p.strip() for p in line.split('|')]
                        if len(parts) >= 2:
                            # Common format: Position | Company | Date
                            # or Company | Position | Date
                            for part in parts:
                                if re.search(r'\d{4}', part):
                                    duration_str = part
                                elif not position and any(kw in part.lower() for kw in ['engineer', 'developer', 'manager', 'analyst', 'designer']):
                                    position = part
                                elif not company:
                                    company = part
                        break
            
            # Strategy 3: Extract from first few lines
            if not position or not company:
                for i, line in enumerate(lines[:4]):
                    line_lower = line.lower()
                    
                    # Check if it's a position (has job title keywords)
                    position_indicators = ['engineer', 'developer', 'manager', 'analyst', 'designer', 'consultant', 
                                          'specialist', 'coordinator', 'director', 'lead', 'senior', 'junior', 
                                          'intern', 'associate', 'assistant', 'head', 'chief']
                    
                    if any(indicator in line_lower for indicator in position_indicators) and not position:
                        position = line
                        continue
                    
                    # Check if it's a company (has company indicators)
                    company_indicators = ['ltd', 'inc', 'corp', 'llc', 'gmbh', 'pt', 'cv', 'technologies', 'solutions', 'group', 'holdings']
                    
                    if (any(indicator in line_lower for indicator in company_indicators) or 
                        (len(line) > 3 and i == 0)) and not company and line != position:
                        company = line
                        continue
            
            # Parse duration
            if duration_str:
                start_date, end_date, years = self._parse_duration(duration_str)
            else:
                # Try to find dates anywhere in the entry
                for line in lines:
                    if re.search(r'\d{4}', line):
                        start_date, end_date, years = self._parse_duration(line)
                        duration_str = line
                        break
            
            # Extract responsibilities and achievements
            responsibilities = []
            achievements = []
            technologies = []
            
            for line in lines:
                # Skip header lines
                if line in [company, position, duration_str]:
                    continue
                
                if len(line) < 10:
                    continue
                
                line_lower = line.lower()
                
                # Achievement indicators (quantifiable results)
                achievement_patterns = [
                    r'\d+%',  # Percentages
                    r'\d+x',  # Multipliers
                    r'increased|improved|reduced|saved|achieved|generated|optimized|enhanced',
                    r'resulting in|led to|contributed to',
                    r'\$\d+|\d+\s*(?:million|thousand|juta|ribu)',  # Money
                    r'\d+\s*(?:users|customers|clients|employees|projects)'  # Numbers
                ]
                
                is_achievement = any(re.search(pattern, line_lower) for pattern in achievement_patterns)
                
                if is_achievement and len(achievements) < 5:
                    achievements.append(line)
                elif len(responsibilities) < 10:
                    # Filter out very generic lines
                    if not any(skip in line_lower for skip in ['responsible for', 'duties include', 'job description']):
                        responsibilities.append(line)
                
                # Extract technologies
                tech_keywords = [
                    'python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'go', 'rust', 'php', 'ruby', 'swift', 'kotlin',
                    'react', 'vue', 'angular', 'node', 'django', 'flask', 'spring', 'laravel', 
                    'sql', 'mysql', 'postgresql', 'mongodb', 'redis', 'elasticsearch',
                    'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'jenkins', 'terraform',
                    'machine learning', 'tensorflow', 'pytorch', 'scikit-learn', 'pandas'
                ]
                
                for tech in tech_keywords:
                    if tech in line_lower and tech.title() not in technologies:
                        technologies.append(tech.title())
            
            # Only add if we have at least company OR position
            if company or position:
                experiences.append(WorkExperience(
                    company=company or "Unknown Company",
                    position=position or "Unknown Position",
                    duration=duration_str or "Unknown Duration",
                    start_date=start_date,
                    end_date=end_date,
                    years=years,
                    responsibilities=responsibilities,
                    achievements=achievements,
                    technologies=technologies,
                    relevance_score=0.0
                ))
        
        return experiences
    
    def _parse_duration(self, duration_str: str) -> Tuple[Optional[str], Optional[str], float]:
        """Parse duration string ke start_date, end_date, years"""
        start_date = None
        end_date = None
        years = 0.0
        
        # Pattern: "Jan 2020 - Dec 2022" or "2020 - 2022" or "2020 - Present"
        date_patterns = [
            r'(\w+\s+\d{4})\s*-\s*(\w+\s+\d{4})',  # Jan 2020 - Dec 2022
            r'(\d{4})\s*-\s*(\d{4})',  # 2020 - 2022
            r'(\d{4})\s*-\s*(present|now|current|sekarang)',  # 2020 - Present
        ]
        
        for pattern in date_patterns:
            match = re.search(pattern, duration_str, re.IGNORECASE)
            if match:
                start_date = match.group(1)
                end_date = match.group(2)
                
                # Calculate years
                try:
                    start_year = int(re.search(r'\d{4}', start_date).group())
                    if re.search(r'\d{4}', end_date):
                        end_year = int(re.search(r'\d{4}', end_date).group())
                    else:
                        end_year = datetime.now().year
                    
                    years = end_year - start_year
                    
                    # Add months if available
                    month_diff = self._calculate_month_difference(start_date, end_date)
                    years += month_diff / 12.0
                    
                except:
                    pass
                
                break
        
        return start_date, end_date, round(years, 1)
    
    def _calculate_month_difference(self, start_str: str, end_str: str) -> int:
        """Calculate month difference between dates"""
        months_map = {
            'jan': 1, 'feb': 2, 'mar': 3, 'apr': 4, 'may': 5, 'jun': 6,
            'jul': 7, 'aug': 8, 'sep': 9, 'oct': 10, 'nov': 11, 'dec': 12
        }
        
        try:
            start_month = 1
            end_month = 1
            
            for month_name, month_num in months_map.items():
                if month_name in start_str.lower():
                    start_month = month_num
                if month_name in end_str.lower():
                    end_month = month_num
            
            return abs(end_month - start_month)
        except:
            return 0
    
    def parse_education(self, education_text: str) -> List[Education]:
        """Parse pendidikan dengan pattern matching yang lebih baik"""
        educations = []
        
        # Split by common patterns
        # Pattern 1: University name on separate line
        # Pattern 2: Degree, Field, Institution in one or multiple lines
        entries = re.split(r'\n(?=[A-Z][a-z]+[\s\w]+(?:University|Institut|Universitas|College|School|Akademi))', education_text)
        
        # If no split, try alternative pattern
        if len(entries) <= 1:
            entries = re.split(r'\n(?=(?:Bachelor|Master|Sarjana|Magister|S1|S2|S3|D3|D4|PhD|Ph\.D))', education_text)
        
        # If still no split, split by years
        if len(entries) <= 1:
            entries = re.split(r'\n(?=\d{4}\s*-\s*\d{4})', education_text)
        
        for entry in entries:
            if len(entry.strip()) < 10:
                continue
            
            lines = [l.strip() for l in entry.split('\n') if l.strip()]
            if not lines:
                continue
            
            institution = ""
            degree = ""
            field_of_study = ""
            graduation_year = None
            gpa = None
            achievements = []
            
            # Strategy 1: Detect institution (usually has University/Institut/College keywords)
            institution_keywords = [
                'university', 'universitas', 'institut', 'college', 'school', 
                'akademi', 'politeknik', 'sekolah tinggi'
            ]
            
            for line in lines:
                line_lower = line.lower()
                if any(kw in line_lower for kw in institution_keywords):
                    institution = line
                    break
            
            # If no institution found, use first line
            if not institution and lines:
                institution = lines[0]
            
            # Strategy 2: Detect degree
            degree_patterns = [
                r'(Bachelor.*?(?:Science|Arts|Engineering|Technology|Computer Science|Information|Business|Economics)?)',
                r'(Master.*?(?:Science|Arts|Engineering|Technology|Computer Science|Information|Business|Economics)?)',
                r'(Sarjana\s+(?:Komputer|Teknik|Ekonomi|Hukum|Sosial|Ilmu)?.*?)',
                r'(Magister.*?)',
                r'(S[123]|D[34])\s*(?:-|–)?\s*([A-Za-z\s]+)?',
                r'(PhD|Ph\.D\.?)\s+(?:in\s+)?([A-Za-z\s]+)?',
                r'(Doctor\s+of\s+Philosophy)\s+(?:in\s+)?([A-Za-z\s]+)?'
            ]
            
            for line in lines:
                for pattern in degree_patterns:
                    match = re.search(pattern, line, re.IGNORECASE)
                    if match:
                        degree = match.group(0).strip()
                        # Try to extract field from the same line
                        field_match = re.search(r'(?:in|of|jurusan|prodi|program studi)\s+([A-Za-z\s&,]+?)(?:\s+\d{4}|\s*$|,|\|)', line, re.IGNORECASE)
                        if field_match:
                            field_of_study = field_match.group(1).strip()
                        break
                if degree:
                    break
            
            # Strategy 3: Extract field of study if not found
            if not field_of_study:
                field_keywords = [
                    'computer science', 'information technology', 'software engineering', 'data science',
                    'business administration', 'economics', 'management', 'accounting', 'finance',
                    'engineering', 'mechanical', 'electrical', 'civil', 'chemical',
                    'teknik informatika', 'sistem informasi', 'ilmu komputer', 'teknik komputer',
                    'manajemen', 'ekonomi', 'akuntansi', 'hukum', 'psikologi'
                ]
                
                text_combined = ' '.join(lines).lower()
                for field_kw in field_keywords:
                    if field_kw in text_combined:
                        field_of_study = field_kw.title()
                        break
            
            # Strategy 4: Extract year
            for line in lines:
                # Pattern: 2015 - 2019, 2015-2019, Class of 2019, Graduated 2019
                year_patterns = [
                    r'(\d{4})\s*-\s*(\d{4})',  # 2015 - 2019
                    r'(?:graduated|lulus|graduated\s+in|class\s+of)\s+(\d{4})',  # Graduated 2019
                    r'(\d{4})\s*$'  # Just year at end of line
                ]
                
                for pattern in year_patterns:
                    match = re.search(pattern, line, re.IGNORECASE)
                    if match:
                        # Take the latest year (graduation year)
                        years = [int(y) for y in match.groups() if y and y.isdigit()]
                        if years:
                            graduation_year = max(years)
                        break
                if graduation_year:
                    break
            
            # Strategy 5: Extract GPA
            for line in lines:
                gpa_patterns = [
                    r'(?:gpa|ipk|cgpa)\s*:?\s*([\d.]+)(?:/[\d.]+)?',
                    r'([\d.]+)\s*/\s*4\.0',
                    r'([\d.]+)\s*/\s*4\.00',
                ]
                
                for pattern in gpa_patterns:
                    match = re.search(pattern, line, re.IGNORECASE)
                    if match:
                        try:
                            gpa_value = float(match.group(1))
                            # Normalize to 4.0 scale if needed
                            if gpa_value > 4.0:
                                gpa = gpa_value / 100 * 4.0  # Assume 100 scale
                            else:
                                gpa = gpa_value
                        except:
                            pass
                        break
                if gpa:
                    break
            
            # Strategy 6: Extract achievements
            achievement_keywords = [
                'cum laude', 'magna cum laude', 'summa cum laude', 'with honors', 'with distinction',
                'dean\'s list', 'dean list', 'scholarship', 'beasiswa', 'penghargaan', 'award',
                'first class', 'second class', 'distinction', 'merit'
            ]
            
            for line in lines:
                line_lower = line.lower()
                for kw in achievement_keywords:
                    if kw in line_lower:
                        achievements.append(line)
                        break
            
            # Build Education object
            educations.append(Education(
                institution=institution or "Unknown Institution",
                degree=degree or "Unknown Degree",
                field_of_study=field_of_study or "Unknown Field",
                graduation_year=graduation_year,
                gpa=gpa,
                achievements=achievements
            ))
        
        return educations if educations else [
            Education(
                institution="Unknown Institution",
                degree="Unknown Degree",
                field_of_study="Unknown Field",
                graduation_year=None,
                gpa=None,
                achievements=[]
            )
        ]
    
    def parse_skills(self, skills_text: str, full_cv_text: str) -> Dict[str, List[str]]:
        """Parse skills ke dalam categories"""
        skills = {
            'technical': [],
            'soft': [],
            'languages': [],
            'tools': []
        }
        
        # Technical skills
        tech_keywords = [
            'python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'php', 'ruby', 'go', 'rust', 'swift', 'kotlin',
            'react', 'vue', 'angular', 'node.js', 'django', 'flask', 'spring', 'laravel', '.net',
            'sql', 'mysql', 'postgresql', 'mongodb', 'redis', 'elasticsearch', 'oracle',
            'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'jenkins', 'terraform',
            'machine learning', 'deep learning', 'ai', 'data science', 'nlp', 'computer vision',
            'tensorflow', 'pytorch', 'keras', 'scikit-learn', 'pandas', 'numpy',
            'html', 'css', 'sass', 'bootstrap', 'tailwind',
            'rest api', 'graphql', 'microservices', 'agile', 'scrum', 'devops', 'ci/cd'
        ]
        
        text_combined = (skills_text + " " + full_cv_text).lower()
        
        for tech in tech_keywords:
            if tech in text_combined and tech not in skills['technical']:
                skills['technical'].append(tech.title())
        
        # Soft skills
        soft_keywords = [
            'leadership', 'communication', 'teamwork', 'problem solving', 'critical thinking',
            'project management', 'time management', 'adaptability', 'creativity',
            'analytical', 'strategic thinking', 'decision making', 'collaboration',
            'kepemimpinan', 'komunikasi', 'kerja tim', 'pemecahan masalah'
        ]
        
        for soft in soft_keywords:
            if soft in text_combined and soft not in skills['soft']:
                skills['soft'].append(soft.title())
        
        # Languages
        language_keywords = [
            'english', 'bahasa indonesia', 'mandarin', 'japanese', 'korean', 'arabic', 'spanish', 'french', 'german',
            'inggris', 'indonesia', 'jepang', 'korea', 'arab', 'spanyol', 'perancis', 'jerman'
        ]
        
        for lang in language_keywords:
            if lang in text_combined and lang not in skills['languages']:
                skills['languages'].append(lang.title())
        
        # Tools
        tool_keywords = [
            'git', 'jira', 'confluence', 'slack', 'trello', 'asana',
            'figma', 'sketch', 'adobe xd', 'photoshop', 'illustrator',
            'tableau', 'power bi', 'excel', 'google analytics',
            'vscode', 'intellij', 'pycharm', 'jupyter'
        ]
        
        for tool in tool_keywords:
            if tool in text_combined and tool not in skills['tools']:
                skills['tools'].append(tool.title())
        
        return skills
    
    def create_cv_profile(self, text: str) -> CVProfile:
        """
        Create comprehensive CV profile
        """
        # TRY ENHANCED PARSER FIRST for maximum extraction
        enhanced_data = self.parse_with_enhanced_parser(text)
        
        if enhanced_data:
            # Use enhanced parser results
            candidate_name = enhanced_data.get('name', 'Unknown')
            
            contact_info = {
                'email': enhanced_data.get('email', ''),
                'phone': enhanced_data.get('phone', ''),
                'linkedin': enhanced_data.get('linkedin', ''),
                'github': enhanced_data.get('github', ''),
                'location': enhanced_data.get('location', ''),
            }
            
            professional_summary = enhanced_data.get('summary', '')
            
            # Convert enhanced experience format to WorkExperience objects
            work_experiences = []
            for exp in enhanced_data.get('experiences', []):
                work_experiences.append(WorkExperience(
                    company=exp.get('company', 'Unknown Company'),
                    position=exp.get('position', 'Unknown Position'),
                    duration=exp.get('duration', ''),
                    start_date=exp.get('start_date'),
                    end_date=exp.get('end_date'),
                    years=exp.get('years', 0),
                    responsibilities=exp.get('responsibilities', []),
                    achievements=exp.get('achievements', []),
                    technologies=exp.get('technologies', []),
                    relevance_score=0.0
                ))
            
            # Convert enhanced education format to Education objects
            education = []
            for edu in enhanced_data.get('education', []):
                education.append(Education(
                    institution=edu.get('institution', 'Unknown Institution'),
                    degree=edu.get('degree', 'Unknown Degree'),
                    field_of_study=edu.get('field', ''),
                    graduation_year=int(edu.get('year')) if edu.get('year') else None,
                    gpa=float(edu.get('gpa')) if edu.get('gpa') else None,
                    achievements=[]
                ))
            
            skills = enhanced_data.get('skills', {
                'technical': [],
                'soft': [],
                'languages': [],
                'tools': []
            })
            
            certifications = enhanced_data.get('certifications', [])
            
            # Convert projects
            projects = []
            for proj in enhanced_data.get('projects', []):
                projects.append({
                    'name': proj.get('name', ''),
                    'description': proj.get('description', '')
                })
            
            awards = enhanced_data.get('awards', [])
            
            # Interests
            interests = enhanced_data.get('interests', [])
            
        else:
            # FALLBACK to original parsing
            # Parse structure
            sections = self.parse_cv_structure(text)
            
            # Extract name with better strategy
            lines = text.split('\n')
            candidate_name = "Unknown"
            
            # Strategy 1: Look in first 15 lines
            for i, line in enumerate(lines[:15]):
                line_clean = line.strip()
                
                # Skip if too short or too long
                if len(line_clean) < 5 or len(line_clean) > 60:
                    continue
                
                # Skip lines with these keywords
                skip_keywords = ['curriculum', 'resume', 'cv', 'vitae', 'phone', 'email', 'address', 
                               'objective', 'summary', 'profile', 'education', 'experience', 'skills',
                               'http', 'www', '@', '.com', '.id', '|']
                
                if any(kw in line_clean.lower() for kw in skip_keywords):
                    continue
                
                # Skip if mostly numbers or symbols
                alpha_count = sum(c.isalpha() for c in line_clean)
                if alpha_count < len(line_clean) * 0.6:
                    continue
                
                # Good indicators for name:
                # - Has at least 2 words
                # - Contains capitals
                # - Not all capitals (unless short like "JOHN DOE")
                words = line_clean.split()
                if len(words) >= 2:
                    has_capitals = any(c.isupper() for c in line_clean)
                    is_all_caps = line_clean == line_clean.upper()
                    
                    # Accept if has capitals and either:
                    # - Not all caps, OR
                    # - All caps but reasonable length (2-4 words)
                    if has_capitals and (not is_all_caps or len(words) <= 4):
                        candidate_name = line_clean
                        break
            
            # Strategy 2: If still unknown, extract from contact or filename
            if candidate_name == "Unknown":
                # Try to find name pattern near email
                for line in lines[:20]:
                    if '@' in line or 'email' in line.lower():
                        # Check previous line
                        idx = lines.index(line)
                        if idx > 0:
                            prev_line = lines[idx-1].strip()
                            if 5 < len(prev_line) < 50 and sum(c.isalpha() for c in prev_line) > 5:
                                candidate_name = prev_line
                                break
            
            # Contact info
            contact_info = self.extract_contact_info(text)
            
            # Professional summary
            professional_summary = ""
            if 'summary' in sections:
                professional_summary = sections['summary'].content[:500]
            
            # Work experiences
            work_experiences = []
            if 'experience' in sections:
                work_experiences = self.parse_work_experiences(sections['experience'].content)
            
            # Education
            education = []
            if 'education' in sections:
                education = self.parse_education(sections['education'].content)
            
            # Skills
            skills_text = sections.get('skills', CVSection('', '', {}, 0)).content
            skills = self.parse_skills(skills_text, text)
            
            # Certifications
            certifications = []
            if 'certifications' in sections:
                cert_lines = sections['certifications'].content.split('\n')
                certifications = [line.strip() for line in cert_lines if len(line.strip()) > 10]
            
            # Projects
            projects = []
            if 'projects' in sections:
                project_entries = sections['projects'].content.split('\n\n')
                for entry in project_entries:
                    if len(entry.strip()) > 20:
                        projects.append({'description': entry.strip()})
            
            # Awards
            awards = []
            if 'awards' in sections:
                award_lines = sections['awards'].content.split('\n')
                awards = [line.strip() for line in award_lines if len(line.strip()) > 10]
            
            # Interests
            interests = []
            if 'interests' in sections:
                interest_text = sections['interests'].content.lower()
                interests = [i.strip() for i in interest_text.split(',') if len(i.strip()) > 3]
        
        # Calculate total experience
        total_experience_years = sum([exp.years for exp in work_experiences])
        
        # Determine seniority level
        if total_experience_years >= 10:
            seniority_level = "Executive/Senior"
        elif total_experience_years >= 5:
            seniority_level = "Senior"
        elif total_experience_years >= 2:
            seniority_level = "Mid-Level"
        else:
            seniority_level = "Entry-Level"
        
        return CVProfile(
            candidate_name=candidate_name,
            contact_info=contact_info,
            professional_summary=professional_summary,
            work_experiences=work_experiences,
            education=education,
            skills=skills,
            certifications=certifications,
            projects=projects,
            awards=awards,
            interests=interests,
            total_experience_years=round(total_experience_years, 1),
            seniority_level=seniority_level
        )


class AdvancedCVAnalyzer:
    """
    Advanced CV Analyzer - Analisis mendalam dengan AI reasoning
    """
    
    def __init__(self, use_mock_models: bool = False):
        self.parser = AdvancedCVParser()
        self.nlp_engine = get_advanced_nlp_engine(use_mock_models=use_mock_models)
        self.reasoning_engine = get_reasoning_engine(self.nlp_engine)
        self.use_mock = use_mock_models
    
    def analyze_cv_deep(
        self,
        cv_text: str,
        criteria: str,
        criteria_details: Optional[Dict[str, Any]] = None
    ) -> CVAnalysisResult:
        """
        Deep analysis CV dengan contextual understanding
        
        Args:
            cv_text: Full CV text
            criteria: Kriteria pencarian/requirements
            criteria_details: Optional detailed criteria breakdown
            
        Returns:
            CVAnalysisResult with comprehensive analysis
        """
        # 1. Create CV Profile
        cv_profile = self.parser.create_cv_profile(cv_text)
        
        # 2. Calculate relevance scores
        relevance_score, relevance_breakdown = self._calculate_deep_relevance(
            cv_profile, criteria
        )
        
        # 3. Identify strengths
        strengths = self._identify_strengths(cv_profile, cv_text)
        
        # 4. Identify potential areas
        potential_areas = self._identify_potential_areas(cv_profile, criteria)
        
        # 5. Fit analysis
        fit_analysis = self._analyze_fit(cv_profile, criteria, cv_text)
        
        # 6. Professional assessment
        professional_assessment = self._generate_professional_assessment(
            cv_profile, relevance_score, strengths, fit_analysis
        )
        
        # 7. Soft skills assessment
        soft_skills_assessment = self._assess_soft_skills(cv_text)
        
        # 8. Overall recommendation
        overall_recommendation = self._generate_recommendation(
            relevance_score, strengths, potential_areas, fit_analysis
        )
        
        # 9. Confidence score
        confidence_score = self._calculate_confidence(cv_profile, cv_text)
        
        # 10. Generate relevance reasoning
        relevance_reasoning = self._generate_relevance_reasoning(
            relevance_breakdown,
            relevance_score,
            criteria
        )
        
        return CVAnalysisResult(
            cv_profile=cv_profile,
            relevance_score=relevance_score,
            relevance_breakdown=relevance_breakdown,
            relevance_reasoning=relevance_reasoning,
            strengths=strengths,
            potential_areas=potential_areas,
            fit_analysis=fit_analysis,
            professional_assessment=professional_assessment,
            soft_skills_assessment=soft_skills_assessment,
            overall_recommendation=overall_recommendation,
            confidence_score=confidence_score
        )
    
    def _calculate_deep_relevance(
        self,
        cv_profile: CVProfile,
        criteria: str
    ) -> Tuple[float, Dict[str, float]]:
        """Calculate relevance dengan breakdown detail"""
        
        breakdown = {}
        
        # 1. Skills match
        all_skills = (
            cv_profile.skills['technical'] +
            cv_profile.skills['soft'] +
            cv_profile.skills['tools']
        )
        skills_text = " ".join(all_skills)
        skills_similarity = self.nlp_engine.calculate_semantic_similarity(
            skills_text, criteria
        )
        breakdown['skills_match'] = skills_similarity * 0.35  # 35% weight
        
        # 2. Experience relevance
        experience_texts = []
        for exp in cv_profile.work_experiences:
            exp_text = f"{exp.position} {exp.company} {' '.join(exp.responsibilities)} {' '.join(exp.achievements)}"
            experience_texts.append(exp_text)
        
        combined_exp = " ".join(experience_texts)
        experience_similarity = self.nlp_engine.calculate_semantic_similarity(
            combined_exp, criteria
        )
        breakdown['experience_match'] = experience_similarity * 0.30  # 30% weight
        
        # 3. Education relevance
        education_texts = [f"{edu.degree} {edu.field_of_study}" for edu in cv_profile.education]
        combined_edu = " ".join(education_texts)
        education_similarity = self.nlp_engine.calculate_semantic_similarity(
            combined_edu, criteria
        )
        breakdown['education_match'] = education_similarity * 0.15  # 15% weight
        
        # 4. Certifications relevance
        cert_text = " ".join(cv_profile.certifications)
        cert_similarity = self.nlp_engine.calculate_semantic_similarity(
            cert_text, criteria
        ) if cert_text else 0.0
        breakdown['certifications_match'] = cert_similarity * 0.10  # 10% weight
        
        # 5. Seniority bonus
        seniority_bonus = 0.0
        if cv_profile.total_experience_years >= 5:
            seniority_bonus = 0.05
        elif cv_profile.total_experience_years >= 3:
            seniority_bonus = 0.03
        breakdown['seniority_bonus'] = seniority_bonus * 0.10  # 10% weight
        
        # Total score
        total_score = sum(breakdown.values())
        total_score = min(1.0, max(0.0, total_score))
        
        return round(total_score, 3), breakdown
    
    def _identify_strengths(self, cv_profile: CVProfile, cv_text: str) -> List[AnalysisInsight]:
        """Identify kekuatan kandidat"""
        strengths = []
        
        # 1. Experience strength
        if cv_profile.total_experience_years >= 5:
            evidence = [
                f"Total {cv_profile.total_experience_years} tahun pengalaman",
                f"Seniority level: {cv_profile.seniority_level}"
            ]
            if cv_profile.work_experiences:
                top_exp = cv_profile.work_experiences[0]
                evidence.append(f"Posisi terakhir: {top_exp.position} di {top_exp.company}")
            
            strengths.append(AnalysisInsight(
                category="Experience",
                score=0.9,
                evidence=evidence,
                reasoning=f"Kandidat memiliki pengalaman profesional yang solid dengan {cv_profile.total_experience_years} tahun di industri.",
                confidence=0.95
            ))
        
        # 2. Technical skills strength
        if len(cv_profile.skills['technical']) >= 5:
            strengths.append(AnalysisInsight(
                category="Technical Skills",
                score=0.85,
                evidence=cv_profile.skills['technical'][:10],
                reasoning=f"Portfolio teknis yang kuat dengan {len(cv_profile.skills['technical'])} keahlian yang teridentifikasi.",
                confidence=0.9
            ))
        
        # 3. Education strength
        if cv_profile.education:
            top_edu = cv_profile.education[0]
            score = 0.7
            if top_edu.gpa and top_edu.gpa >= 3.5:
                score = 0.9
            
            evidence = [f"{top_edu.degree} - {top_edu.field_of_study}"]
            if top_edu.gpa:
                evidence.append(f"GPA: {top_edu.gpa}")
            if top_edu.achievements:
                evidence.extend(top_edu.achievements[:2])
            
            strengths.append(AnalysisInsight(
                category="Education",
                score=score,
                evidence=evidence,
                reasoning=f"Background pendidikan yang relevan dari {top_edu.institution}.",
                confidence=0.85
            ))
        
        # 4. Leadership/Impact
        leadership_indicators = 0
        impact_evidence = []
        for exp in cv_profile.work_experiences:
            if exp.achievements:
                leadership_indicators += len(exp.achievements)
                impact_evidence.extend(exp.achievements[:2])
        
        if leadership_indicators >= 3:
            strengths.append(AnalysisInsight(
                category="Leadership & Impact",
                score=0.85,
                evidence=impact_evidence[:5],
                reasoning=f"Menunjukkan {leadership_indicators} bukti pencapaian dan dampak yang terukur.",
                confidence=0.8
            ))
        
        # 5. Certifications
        if len(cv_profile.certifications) >= 2:
            strengths.append(AnalysisInsight(
                category="Certifications",
                score=0.75,
                evidence=cv_profile.certifications[:5],
                reasoning=f"Memiliki {len(cv_profile.certifications)} sertifikasi profesional yang menunjukkan komitmen pada pengembangan diri.",
                confidence=0.85
            ))
        
        return strengths
    
    def _identify_potential_areas(
        self,
        cv_profile: CVProfile,
        criteria: str
    ) -> List[AnalysisInsight]:
        """Identify area potensial improvement"""
        potential_areas = []
        
        # Extract entities from criteria
        criteria_entities = self.nlp_engine.extract_entities(criteria)
        
        # Check missing skills
        criteria_skills = criteria_entities.get('skills', [])
        cv_skills_flat = (
            cv_profile.skills['technical'] +
            cv_profile.skills['tools']
        )
        
        missing_skills = []
        for skill in criteria_skills:
            if not any(skill.lower() in cv_skill.lower() for cv_skill in cv_skills_flat):
                missing_skills.append(skill)
        
        if missing_skills and len(missing_skills) <= 5:
            potential_areas.append(AnalysisInsight(
                category="Skills Gap",
                score=0.6,
                evidence=missing_skills,
                reasoning=f"Beberapa skills yang disebutkan dalam kriteria tidak terdeteksi dalam CV: {', '.join(missing_skills[:3])}",
                confidence=0.7
            ))
        
        # Check experience level
        if criteria_entities.get('education'):
            if not cv_profile.education:
                potential_areas.append(AnalysisInsight(
                    category="Education",
                    score=0.5,
                    evidence=["Informasi pendidikan tidak lengkap atau tidak ditemukan"],
                    reasoning="Kriteria menyebutkan requirements pendidikan, namun informasi pendidikan dalam CV kurang detail.",
                    confidence=0.6
                ))
        
        return potential_areas
    
    def _analyze_fit(
        self,
        cv_profile: CVProfile,
        criteria: str,
        cv_text: str
    ) -> Dict[str, Any]:
        """Analyze overall fit"""
        
        # Experience depth analysis
        exp_analysis = self.reasoning_engine.analyze_experience_depth(cv_text)
        
        # Extract topics from CV
        cv_topics = self.nlp_engine.extract_semantic_topics(cv_text, num_topics=5)
        
        # Extract topics from criteria
        criteria_topics = self.nlp_engine.extract_semantic_topics(criteria, num_topics=5)
        
        # Topic overlap
        topic_overlap = self._calculate_topic_overlap(cv_topics, criteria_topics)
        
        fit_analysis = {
            'experience_depth': exp_analysis,
            'cv_topics': [topic for topic, score in cv_topics],
            'criteria_topics': [topic for topic, score in criteria_topics],
            'topic_overlap': topic_overlap,
            'seniority_match': self._assess_seniority_match(cv_profile, criteria),
            'cultural_indicators': self._extract_cultural_indicators(cv_text)
        }
        
        return fit_analysis
    
    def _calculate_topic_overlap(
        self,
        cv_topics: List[Tuple[str, float]],
        criteria_topics: List[Tuple[str, float]]
    ) -> float:
        """Calculate topic overlap score"""
        if not cv_topics or not criteria_topics:
            return 0.0
        
        cv_topic_strs = [topic.lower() for topic, _ in cv_topics]
        criteria_topic_strs = [topic.lower() for topic, _ in criteria_topics]
        
        # Count semantic matches
        matches = 0
        for cv_topic in cv_topic_strs:
            for crit_topic in criteria_topic_strs:
                # Simple word overlap check
                cv_words = set(cv_topic.split())
                crit_words = set(crit_topic.split())
                overlap = len(cv_words & crit_words)
                if overlap >= 1:
                    matches += 1
                    break
        
        overlap_score = matches / len(criteria_topic_strs) if criteria_topic_strs else 0.0
        return round(overlap_score, 2)
    
    def _assess_seniority_match(self, cv_profile: CVProfile, criteria: str) -> str:
        """Assess if seniority matches criteria"""
        criteria_lower = criteria.lower()
        
        cv_level = cv_profile.seniority_level.lower()
        
        if any(word in criteria_lower for word in ['senior', 'lead', 'principal', 'staff']):
            required_level = "senior"
        elif any(word in criteria_lower for word in ['junior', 'entry', 'fresh graduate']):
            required_level = "entry"
        else:
            required_level = "mid"
        
        if required_level in cv_level or cv_level in required_level:
            return "Excellent Match"
        elif (required_level == "senior" and "mid" in cv_level):
            return "Good Match (slightly below)"
        elif (required_level == "entry" and "mid" in cv_level):
            return "Over-qualified"
        else:
            return "Partial Match"
    
    def _extract_cultural_indicators(self, cv_text: str) -> List[str]:
        """Extract cultural fit indicators"""
        indicators = []
        
        text_lower = cv_text.lower()
        
        cultural_keywords = {
            'team_player': ['team', 'collaboration', 'collaborative', 'cooperative'],
            'innovative': ['innovative', 'creative', 'innovation', 'creativity'],
            'results_driven': ['achieved', 'delivered', 'results', 'success', 'accomplished'],
            'adaptable': ['adapt', 'flexible', 'versatile', 'agile'],
            'continuous_learner': ['learning', 'training', 'development', 'course', 'certification']
        }
        
        for trait, keywords in cultural_keywords.items():
            if sum(1 for kw in keywords if kw in text_lower) >= 2:
                indicators.append(trait.replace('_', ' ').title())
        
        return indicators
    
    def _generate_professional_assessment(
        self,
        cv_profile: CVProfile,
        relevance_score: float,
        strengths: List[AnalysisInsight],
        fit_analysis: Dict[str, Any]
    ) -> str:
        """Generate professional HR-style assessment"""
        
        assessment_parts = []
        
        # Opening
        assessment_parts.append(
            f"**Kandidat: {cv_profile.candidate_name}**\n\n"
            f"Seniority: {cv_profile.seniority_level} | "
            f"Total Experience: {cv_profile.total_experience_years} tahun | "
            f"Relevance Score: {relevance_score:.1%}\n"
        )
        
        # Profile summary
        assessment_parts.append(
            f"\n**Ringkasan Profil:**\n"
            f"Kandidat dengan background profesional di "
        )
        
        if cv_profile.work_experiences:
            recent_exp = cv_profile.work_experiences[0]
            assessment_parts.append(
                f"{recent_exp.position} dengan pengalaman di {recent_exp.company}. "
            )
        
        if cv_profile.education:
            top_edu = cv_profile.education[0]
            assessment_parts.append(
                f"Lulusan {top_edu.degree} dari {top_edu.institution}. "
            )
        
        # Key strengths
        assessment_parts.append(
            f"\n\n**Kekuatan Utama:**\n"
        )
        for i, strength in enumerate(strengths[:3], 1):
            assessment_parts.append(
                f"{i}. **{strength.category}**: {strength.reasoning}\n"
            )
        
        # Experience depth
        exp_depth = fit_analysis.get('experience_depth', {})
        if exp_depth:
            assessment_parts.append(
                f"\n**Analisis Pengalaman:**\n"
                f"- Level: {exp_depth.get('responsibility_level', 'N/A').title()}\n"
                f"- Technical Depth: {exp_depth.get('technical_depth', 'N/A').title()}\n"
            )
            
            if exp_depth.get('leadership_indicators'):
                assessment_parts.append(
                    f"- Leadership: {len(exp_depth['leadership_indicators'])} indikator ditemukan\n"
                )
        
        # Cultural fit
        cultural_indicators = fit_analysis.get('cultural_indicators', [])
        if cultural_indicators:
            assessment_parts.append(
                f"\n**Indikator Cultural Fit:**\n"
                f"{', '.join(cultural_indicators)}\n"
            )
        
        return "".join(assessment_parts)
    
    def _assess_soft_skills(self, cv_text: str) -> Dict[str, float]:
        """Assess soft skills dari CV text"""
        
        soft_skills = {
            'leadership': 0.0,
            'communication': 0.0,
            'problem_solving': 0.0,
            'teamwork': 0.0,
            'adaptability': 0.0
        }
        
        text_lower = cv_text.lower()
        
        # Leadership
        leadership_kw = ['lead', 'manage', 'supervise', 'mentor', 'coach', 'direct']
        soft_skills['leadership'] = min(1.0, sum(text_lower.count(kw) for kw in leadership_kw) / 10)
        
        # Communication
        comm_kw = ['present', 'communicate', 'speak', 'write', 'articulate', 'explain']
        soft_skills['communication'] = min(1.0, sum(text_lower.count(kw) for kw in comm_kw) / 10)
        
        # Problem solving
        problem_kw = ['solve', 'problem', 'troubleshoot', 'debug', 'resolve', 'fix']
        soft_skills['problem_solving'] = min(1.0, sum(text_lower.count(kw) for kw in problem_kw) / 10)
        
        # Teamwork
        team_kw = ['team', 'collaborate', 'cooperate', 'together', 'group']
        soft_skills['teamwork'] = min(1.0, sum(text_lower.count(kw) for kw in team_kw) / 10)
        
        # Adaptability
        adapt_kw = ['adapt', 'flexible', 'versatile', 'learn', 'change']
        soft_skills['adaptability'] = min(1.0, sum(text_lower.count(kw) for kw in adapt_kw) / 10)
        
        return {k: round(v, 2) for k, v in soft_skills.items()}
    
    def _generate_recommendation(
        self,
        relevance_score: float,
        strengths: List[AnalysisInsight],
        potential_areas: List[AnalysisInsight],
        fit_analysis: Dict[str, Any]
    ) -> str:
        """Generate overall recommendation"""
        
        if relevance_score >= 0.75:
            recommendation = "**HIGHLY RECOMMENDED** ✅\n\n"
            recommendation += "Kandidat menunjukkan kesesuaian yang sangat tinggi dengan kriteria. "
        elif relevance_score >= 0.6:
            recommendation = "**RECOMMENDED** 👍\n\n"
            recommendation += "Kandidat memiliki kesesuaian yang baik dengan sebagian besar kriteria. "
        elif relevance_score >= 0.4:
            recommendation = "**CONSIDER WITH CAUTION** ⚠️\n\n"
            recommendation += "Kandidat memiliki potensi namun ada beberapa gap yang perlu dipertimbangkan. "
        else:
            recommendation = "**NOT RECOMMENDED** ❌\n\n"
            recommendation += "Kesesuaian dengan kriteria masih terbatas. "
        
        # Add key points
        if strengths:
            recommendation += f"Memiliki {len(strengths)} kekuatan utama yang menonjol. "
        
        if potential_areas:
            recommendation += f"Ada {len(potential_areas)} area yang perlu dikembangkan. "
        
        # Seniority match
        seniority_match = fit_analysis.get('seniority_match', '')
        if seniority_match:
            recommendation += f"\n\nSeniority Level: {seniority_match}"
        
        return recommendation
    
    def _calculate_confidence(self, cv_profile: CVProfile, cv_text: str) -> float:
        """Calculate confidence score of analysis"""
        
        confidence = 0.5  # Base confidence
        
        # Increase confidence based on data completeness
        if cv_profile.contact_info.get('email'):
            confidence += 0.05
        
        if cv_profile.work_experiences:
            confidence += 0.15
        
        if cv_profile.education:
            confidence += 0.1
        
        if len(cv_profile.skills['technical']) >= 5:
            confidence += 0.1
        
        if cv_profile.certifications:
            confidence += 0.05
        
        if len(cv_text) > 1000:
            confidence += 0.05
        
        return min(1.0, confidence)
    
    def _generate_relevance_reasoning(
        self,
        breakdown: Dict[str, float],
        overall_score: float,
        criteria: str
    ) -> str:
        """Generate natural language reasoning for relevance score"""
        
        # Sort breakdown by score
        sorted_breakdown = sorted(breakdown.items(), key=lambda x: x[1], reverse=True)
        
        reasoning_parts = []
        
        # Overall assessment
        if overall_score >= 0.8:
            reasoning_parts.append(f"Sangat relevan dengan kriteria '{criteria}'.")
        elif overall_score >= 0.6:
            reasoning_parts.append(f"Cukup relevan dengan kriteria '{criteria}'.")
        elif overall_score >= 0.4:
            reasoning_parts.append(f"Relevansi sedang dengan kriteria '{criteria}'.")
        else:
            reasoning_parts.append(f"Relevansi rendah dengan kriteria '{criteria}'.")
        
        # Top strengths
        top_3 = sorted_breakdown[:3]
        if top_3:
            reasoning_parts.append("\n\n**Komponen Relevansi Tertinggi:**")
            for component, score in top_3:
                if score >= 0.7:
                    level = "Sangat baik"
                elif score >= 0.5:
                    level = "Baik"
                elif score >= 0.3:
                    level = "Cukup"
                else:
                    level = "Kurang"
                
                component_name = component.replace('_', ' ').title()
                reasoning_parts.append(f"- {component_name}: {level} ({score:.1%})")
        
        # Weak areas
        weak_areas = [(k, v) for k, v in sorted_breakdown if v < 0.3]
        if weak_areas:
            reasoning_parts.append("\n\n**Area yang Perlu Ditingkatkan:**")
            for component, score in weak_areas[:2]:
                component_name = component.replace('_', ' ').title()
                reasoning_parts.append(f"- {component_name} ({score:.1%})")
        
        return "\n".join(reasoning_parts)


def get_advanced_cv_analyzer(use_mock_models: bool = False) -> AdvancedCVAnalyzer:
    """Get advanced CV analyzer instance"""
    return AdvancedCVAnalyzer(use_mock_models=use_mock_models)
