"""
CV Parser - Universal document parser untuk CV/Resume
Mendukung PDF, DOCX, TXT, dan ekstraksi text berkualitas tinggi
"""

import os
import io
from typing import Dict, Optional, List, Tuple
from pathlib import Path
import re


class CVParser:
    """
    Universal CV parser dengan dukungan multiple formats
    
    Supported formats:
    - PDF (.pdf)
    - Word Document (.docx)
    - Plain Text (.txt)
    """
    
    def __init__(self):
        self.supported_extensions = ['.pdf', '.docx', '.txt']
    
    def parse(self, file_path_or_bytes, filename: Optional[str] = None) -> Dict[str, any]:
        """
        Parse CV file dan extract text + metadata
        
        Args:
            file_path_or_bytes: File path (str/Path) atau file bytes
            filename: Nama file (required jika input adalah bytes)
        
        Returns:
            Dict dengan keys: text, metadata, sections, success
        """
        try:
            # Determine file type
            if isinstance(file_path_or_bytes, (str, Path)):
                file_path = Path(file_path_or_bytes)
                filename = file_path.name
                extension = file_path.suffix.lower()
                
                with open(file_path, 'rb') as f:
                    file_bytes = f.read()
            else:
                # Bytes input
                file_bytes = file_path_or_bytes
                if not filename:
                    raise ValueError("Filename required when input is bytes")
                
                extension = Path(filename).suffix.lower()
            
            # Validate extension
            if extension not in self.supported_extensions:
                return {
                    "success": False,
                    "error": f"Unsupported file type: {extension}",
                    "text": "",
                    "metadata": {},
                    "sections": {}
                }
            
            # Parse based on file type
            if extension == '.pdf':
                text, metadata = self._parse_pdf(file_bytes)
            elif extension == '.docx':
                text, metadata = self._parse_docx(file_bytes)
            elif extension == '.txt':
                text, metadata = self._parse_txt(file_bytes)
            else:
                text, metadata = "", {}
            
            # Extract sections
            sections = self._extract_sections(text)
            
            # Clean text
            text_cleaned = self._clean_text(text)
            
            return {
                "success": True,
                "text": text_cleaned,
                "raw_text": text,
                "metadata": {
                    **metadata,
                    "filename": filename,
                    "file_type": extension,
                    "text_length": len(text_cleaned),
                    "word_count": len(text_cleaned.split())
                },
                "sections": sections,
                "error": None
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": str(e),
                "text": "",
                "metadata": {"filename": filename if filename else "unknown"},
                "sections": {}
            }
    
    def _parse_pdf(self, pdf_bytes: bytes) -> Tuple[str, Dict]:
        """Parse PDF file"""
        try:
            import PyPDF2
            
            pdf_file = io.BytesIO(pdf_bytes)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text_parts = []
            for page in pdf_reader.pages:
                text_parts.append(page.extract_text())
            
            text = "\n".join(text_parts)
            
            metadata = {
                "pages": len(pdf_reader.pages),
                "parser": "PyPDF2"
            }
            
            # Try to get PDF metadata
            if pdf_reader.metadata:
                metadata.update({
                    "title": pdf_reader.metadata.get('/Title', ''),
                    "author": pdf_reader.metadata.get('/Author', ''),
                    "creator": pdf_reader.metadata.get('/Creator', '')
                })
            
            return text, metadata
            
        except ImportError:
            # Fallback to pdfplumber
            try:
                import pdfplumber
                
                pdf_file = io.BytesIO(pdf_bytes)
                text_parts = []
                
                with pdfplumber.open(pdf_file) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)
                    
                    text = "\n".join(text_parts)
                    metadata = {
                        "pages": len(pdf.pages),
                        "parser": "pdfplumber"
                    }
                
                return text, metadata
                
            except ImportError:
                raise ImportError(
                    "PDF parsing requires PyPDF2 or pdfplumber. "
                    "Install with: pip install PyPDF2 or pip install pdfplumber"
                )
    
    def _parse_docx(self, docx_bytes: bytes) -> Tuple[str, Dict]:
        """Parse DOCX file"""
        try:
            from docx import Document
            
            docx_file = io.BytesIO(docx_bytes)
            doc = Document(docx_file)
            
            # Extract paragraphs
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        text_parts.append(row_text)
            
            text = "\n".join(text_parts)
            
            metadata = {
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "parser": "python-docx"
            }
            
            # Try to get doc properties
            try:
                core_props = doc.core_properties
                metadata.update({
                    "title": core_props.title or "",
                    "author": core_props.author or "",
                    "created": str(core_props.created) if core_props.created else ""
                })
            except:
                pass
            
            return text, metadata
            
        except ImportError:
            raise ImportError(
                "DOCX parsing requires python-docx. "
                "Install with: pip install python-docx"
            )
    
    def _parse_txt(self, txt_bytes: bytes) -> Tuple[str, Dict]:
        """Parse plain text file"""
        # Try different encodings
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                text = txt_bytes.decode(encoding)
                metadata = {
                    "encoding": encoding,
                    "parser": "text"
                }
                return text, metadata
            except UnicodeDecodeError:
                continue
        
        # Fallback: decode with errors ignored
        text = txt_bytes.decode('utf-8', errors='ignore')
        metadata = {
            "encoding": "utf-8 (with errors ignored)",
            "parser": "text"
        }
        return text, metadata
    
    def _clean_text(self, text: str) -> str:
        """Clean dan normalize text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove page numbers (common patterns)
        text = re.sub(r'Page \d+', '', text, flags=re.IGNORECASE)
        text = re.sub(r'\d+\s*/\s*\d+', '', text)
        
        # Clean up
        text = text.strip()
        
        return text
    
    def _extract_sections(self, text: str) -> Dict[str, str]:
        """
        Extract CV sections (education, experience, skills, etc.)
        
        Returns:
            Dict dengan section names sebagai keys
        """
        if not text:
            return {}
        
        sections = {}
        
        # Common section headers (English + Indonesian)
        section_patterns = {
            'contact': r'(?:contact|kontak|alamat|address|email|phone|telepon)',
            'summary': r'(?:summary|profile|ringkasan|profil|objective|tentang)',
            'experience': r'(?:experience|pengalaman|work history|riwayat kerja|employment)',
            'education': r'(?:education|pendidikan|academic|akademik|qualification)',
            'skills': r'(?:skills|keahlian|kemampuan|competencies|kompetensi)',
            'projects': r'(?:projects|proyek|portfolio|works)',
            'certifications': r'(?:certifications?|sertifikat|licenses|lisensi)',
            'languages': r'(?:languages?|bahasa)',
            'awards': r'(?:awards?|achievements?|penghargaan|prestasi)',
            'references': r'(?:references?|referensi)'
        }
        
        # Try to split text by section headers
        lines = text.split('\n')
        current_section = 'header'
        section_content = []
        
        for line in lines:
            line_lower = line.lower().strip()
            
            # Check if line is a section header
            matched_section = None
            for section_name, pattern in section_patterns.items():
                if re.search(pattern, line_lower) and len(line.strip()) < 50:
                    matched_section = section_name
                    break
            
            if matched_section:
                # Save previous section
                if section_content:
                    sections[current_section] = '\n'.join(section_content).strip()
                
                # Start new section
                current_section = matched_section
                section_content = []
            else:
                # Add to current section
                if line.strip():
                    section_content.append(line.strip())
        
        # Save last section
        if section_content:
            sections[current_section] = '\n'.join(section_content).strip()
        
        return sections
    
    def extract_contact_info(self, text: str) -> Dict[str, str]:
        """Extract contact information dari CV text"""
        contact = {}
        
        # Email
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        if emails:
            contact['email'] = emails[0]
        
        # Phone (various formats)
        phone_pattern = r'(?:\+\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3,4}[-.\s]?\d{4}'
        phones = re.findall(phone_pattern, text)
        if phones:
            contact['phone'] = phones[0]
        
        # LinkedIn
        linkedin_pattern = r'(?:linkedin\.com/in/|linkedin\.com/pub/)([A-Za-z0-9_-]+)'
        linkedin = re.search(linkedin_pattern, text, re.IGNORECASE)
        if linkedin:
            contact['linkedin'] = f"linkedin.com/in/{linkedin.group(1)}"
        
        # GitHub
        github_pattern = r'(?:github\.com/)([A-Za-z0-9_-]+)'
        github = re.search(github_pattern, text, re.IGNORECASE)
        if github:
            contact['github'] = f"github.com/{github.group(1)}"
        
        return contact
    
    def extract_keywords(self, text: str, top_n: int = 20) -> List[str]:
        """Extract important keywords dari CV"""
        if not text:
            return []
        
        # Simple keyword extraction using frequency
        # Remove common words
        common_words = set([
            'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for',
            'of', 'with', 'by', 'from', 'as', 'is', 'was', 'are', 'were', 'been',
            'be', 'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would',
            'yang', 'dan', 'atau', 'dengan', 'untuk', 'dari', 'ke', 'di', 'pada',
            'adalah', 'ini', 'itu', 'saya', 'anda', 'telah', 'sudah', 'akan'
        ])
        
        # Tokenize and count
        words = re.findall(r'\b[a-zA-Z]{3,}\b', text.lower())
        word_freq = {}
        
        for word in words:
            if word not in common_words:
                word_freq[word] = word_freq.get(word, 0) + 1
        
        # Get top keywords
        sorted_words = sorted(word_freq.items(), key=lambda x: x[1], reverse=True)
        keywords = [word for word, freq in sorted_words[:top_n]]
        
        return keywords


# Helper functions
def parse_cv(file_path_or_bytes, filename: Optional[str] = None) -> Dict:
    """Quick parse CV file"""
    parser = CVParser()
    return parser.parse(file_path_or_bytes, filename)


def extract_cv_text(file_path_or_bytes, filename: Optional[str] = None) -> str:
    """Extract just the text from CV"""
    result = parse_cv(file_path_or_bytes, filename)
    return result.get('text', '') if result['success'] else ''


if __name__ == "__main__":
    # Example usage
    print("📄 CV Parser - Testing")
    print("=" * 60)
    
    # Test with a sample text
    sample_cv = """
    JOHN DOE
    Senior Software Engineer
    
    Contact:
    Email: john.doe@example.com
    Phone: +1-555-0123
    LinkedIn: linkedin.com/in/johndoe
    
    PROFESSIONAL SUMMARY
    Experienced software engineer with 8+ years in full-stack development.
    
    EXPERIENCE
    Senior Software Engineer | Tech Corp | 2020-Present
    - Led development of microservices architecture
    - Mentored junior developers
    
    EDUCATION
    Bachelor of Computer Science | University of Technology | 2015
    
    SKILLS
    Python, Java, React, Docker, Kubernetes, AWS
    """
    
    # Create sample file
    test_file = Path("test_cv.txt")
    test_file.write_text(sample_cv)
    
    # Parse
    parser = CVParser()
    result = parser.parse(test_file)
    
    print(f"✅ Success: {result['success']}")
    print(f"📝 Text length: {result['metadata']['text_length']} chars")
    print(f"📊 Word count: {result['metadata']['word_count']} words")
    print(f"\n📑 Sections found: {list(result['sections'].keys())}")
    
    # Extract contact
    contact = parser.extract_contact_info(sample_cv)
    print(f"\n📞 Contact info: {contact}")
    
    # Extract keywords
    keywords = parser.extract_keywords(sample_cv, top_n=10)
    print(f"\n🔑 Keywords: {', '.join(keywords)}")
    
    # Cleanup
    test_file.unlink()
    
    print("\n✅ CV Parser test completed!")
